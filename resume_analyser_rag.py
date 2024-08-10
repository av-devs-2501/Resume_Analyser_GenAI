# Import libraries (global) - Start ----------------------------
import streamlit as st
from dotenv import load_dotenv
import uuid
import os
import win32com.client
from pypdf import PdfReader
import docx

# Import libraries (global) - End ------------------------------------------

### All User Defined Functions ------- Start--------------------------------

# --------------Job Descriotion Text Extraction (Start)---------------------

def get_jd_text(job_description):
    # Create a Document object
    doc = docx.Document(job_description)

    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

        #Extract text from tables
        for table in doc.tables:
            for rows in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

    jd = "\\n".join(full_text)

    return jd

# from langchain.schema import Document
#iterate over files in
#that user uploaded PDF files, one by one
def extract_job_description_text(filename, unique_id):
    from langchain.schema import Document
    docs = []
    chunk =  get_jd_text(filename)
    # Adding items to our list - Adding data & its metadata
    docs.append(Document(
        page_content= chunks,
        metadata = {"name": filename.name, "unique_id": unique_id}
    ))
    return docs

# ------------------ Job description Text Extraction (End)---------------------

### --------------------Resume Text Extraction (Start) ------------------------
# Extract text from PDF, DOCX and DOC

## --------------------- PDF Text Extraction (Start) --------------------------
# Extract information from PDF file
def get_pdf_text(pdf_doc):
    text = ""
    pdf_reader = PdfReader(pdf_doc)
    for oage in pdf_reader.pages:
        text += page.extract_text()

    return text

# interates over files in user uploaded PDF files one by one
def create_docs_from_pdf(user_pdf_list, unique_id):
    from langchain.schema import Document
    docs = []
    for filename in user_pdf_list:
        chunk = get_pdf_text(filename)
        #Adding items to our list - Adding data & its metadata
        docs.append(Document(
            page_content= chunks,
            metadata = {"name": filename.name, "unique_id": unique_id}
        ))

    return docs

## --------------------- PDF Text Extraction (End) --------------------------

## --------------------- DOCX Text Extraction (Start) -----------------------

# Extract Information from DOCX file
import docx

def get_docx_text(docx_doc):
    text = ""
    # LOAD the .docx file
    from docx import Document
    doc = Document(docx_doc)
    # Iterate through each paragraph in the document
    for para in doc.paragraph:
        #Add the text from each paragraph to the text variable
        text += para.text + '\n'
    return text



# interates over files in user uploaded DOCX files one by one
def create_docs_from_docx(user_docx_list, unique_id):
    from langchain.schema import Document
    docs = []
    for filename in user_docx_list:
        chunks = get_docx_text(filename)
        docs.append(Document(
            page_content = chunks,
            metadata = {"name": filename.name, "unique_id": unique_id}
        ))
    return docs

## --------------------- DOCX Text Extraction (End) -------------------------


## --------------------- DOC Text Extraction (Start) ------------------------


import comtypes.client
import os

def get_doc_content(doc_file_path):
    #Initialize COM process
    comtypes.CoInitialize()
    try:
        # Load Word:
        word = comtypes.client.CreateObject('Word.Application')
        #Open the document
        doc = word.Document.Open(doc_file_path)
        # Read the content
        content = doc.Range().Text
        # Close the document
        doc.Close(False)
        return content
    finally:
        # Quit wordl application:
        word.Quit()
        # Uninitialize COM process
        comtypes.CoUninitialize()

def save_and_extract_doc_text(file):
    from langchain.schema import Document
    if file:
        doc_contents = []
        filename = "temp.doc" # Customize the filename as needed
        # Specify the directory path:
        save_path = "c:\\User\\AV_DEVS\\Documents\\Gen AI\\Sample Resume for Training_Jan03, 2024\\DOC\\" + filename
        with open(save_path, "wb") as f:
            f.write(file.read())
        # Get the content of the doc file and add it to the list
        doc_contents = (get_doc_content(save_path))
        os.remove(save_path)
    return doc_contents


# interates over files in user uploaded DOC files one by one

def create_docs_from_doc(user_doc_list, unique_id):
    from langchain.schema import Document
    docs = []
    for filename in user_doc_list:
        chunks = save_and_extract_doc_text(filename)
        docs.append(Document(
            page_content = chunks,
            metadata = {"name": filename.name, "unique_id": unique_id}
        ))
    return docs

## --------------------- DOC Text Extraction (End) ------------------------

## -------------------Push Embedding to FAISS (Start) ---------------------

# Import main libraries:
import openai
import os

#Open AI Environment setup:
openai.api.type = "azure"
openai.api.version = "2023-07-01-preview"
openai.api.base = "http://avdevsitrain.openai.azure.com/"
openai.api.key = "ce215cd2444404105b259a5eae769cef2"

#Initializing OpenAI Embeddings:
from langchain_community.embeddings import OpenAIEmbeddings
embeddings = OpenAIEmbeddings (openai_api_key = "ce215cd245804105b259a5eae769cef2", engine 'textembeddingada002')
print(embeddings)

#Create Vector DB and run Similarity Search:
from langchain_community.vectorstores import FAISSdb_faiss = FAISS.from_documents(final_docs_list, embeddings)
results = db_faiss.similarity_search(job_description_text, k = result_count, return_score = True)
return results

## -------------------Push Embedding to FAISS (End) ---------------------

## ------------------- Resume Summerization (Start) ---------------------

def generate_summery(resume_text):
    #using OpenAI's Completion module that helps perform text manipulation

    import openai
    import os
    
    prompt = '''
    Please highlight candidate name, contact number from the following resume and then summarize the resume under three bullet popint headers:
    Work Experiemce, Educational Qualification, Technical Skills.
    '''

    # Set up OpenAI API key:
    openai.api.type = "azure"
    openai.api.version = "2023-07-01-preview"
    openai.api.base = "http://avdevsitrain.openai.azure.com/"
    openai.api.key = "ce215cd2444404105b259a5eae769cef2"

    # Resume Summarization:
    response = opena.ChatCompletion.create(
        engine = "gpt35turbo",
        messages = [{"role": "system", "content": prompt},
        {"role": "user", "content": resume_text}],
        temperature = 0.2,
        max_tokens = 4000,
        top_p = 0.95,
        frequency_penalty = 0,
        presence_penalty = 0,
        stop = None
    )
    return response["choices"][0]["message"]["content"]

## --------------------------- Resume Summaruzation (End) ----------------------

### All User Defined Functions (END) -------------------------------------------

### Main Application in Streamlit (start) --------------------------------------

# Creating a session variable to assign a unique value for each instance of usage:
if 'unique_id' not st.session_state:
    st.session_state['unique_id'] = ''

## -------------------- Streamlit app function (Start) ---------------------------

def main():
    load_dotenv()

    # Application Header:

    st.set_page_config(page_title = "Resume Screening Assistant")
    st.title("Resume Screening Assistant 🤖")
    st.subheader("I can help you in revenue screening process")

    # Upload the JD (DOCX file):
    job_description_file = st.file_uploader("Upload the JD here", type = ["PDF", "DOCX"])

    # Enter count of resumes to be shown in result:
    default_count = '3'
    result_count = int(st.text_input("No. of resumes to return", value - default_count, key - "2"))

    # Upload the Resumes (PDF, DOCX and DOC files):
    resume_files = st.file_uploader("Upload resumes here", type = ["PDF", "DOCX", "DOC"], accept_multiple_files = True)

    # Submit button to steart analysing resume:
    submit = st.button("Analyse resume 🔎")

    # Distributing the resume files in three different variables basied on the ir files types:
    pdf_files = [f for f in resume_files if f.name.endswith('.pdf')]
    docx_files = [f for f in resume_files if f.name.endswith('.docx')]
    doc_files = [f for f in resume_files if f.name.endswith('.doc')]

    if submit:
        with st.spinner('Wait for it...'):

            # Creating a unique ID, so that we can use to query and get only the user uploaded documents form FIASS vector store
            st.session_state['unique_id'] = uuid.uuid4().hex

            # Extracting content from Job Description document using UDF:
            if job_description_file.name.endswith('.pdf'):
                job_description = get_pdf_text(job_description_file) # For PDF
            else:
                job_description = get_docx_text(job_description_file) # For DOCX

            # Create a document list out of all the user uploaded pdf files:
            final_docs_list = create_docs_from_pdf(pdf_files, st.session_state['unique_id'])
            final_docs_list += create_docs_from_docx(docx_files, st.session_state['unique_id'])
            final_docs_list += create_docs_from_doc(doc_files, st.session_state['unique_id'])

            # Similarity Search using FAISS:
            relevant_results = push_to_FAISS_for_similarity_search(job_description, final_docs_list, result_count)

            # Print top resumes and resume summery:
            rank = 0
            st.write("**Here are the top resumes fitting the job description you have uploaded: **")
            for each in relevant_results:
                rank += 1
                st.write("**Rank**:", rank) # adds ranks to each matching resume
                st.write("**Resume File Name:**")
                st.write(each.metadata["name"]) # Shows top matching resume file names
                resume_summery = generate_summery(each.page_content)
                st.write("**Resume Summary:**")
                st.write(resume_summery)
                st.divider() #divider line separating each result

        st.success("Top results presented")
## ------------------------------------Streamlit app function (End)--------------------------------------------------------------------

#Invoking main function:
if __name__== '__main__':
    main()





















